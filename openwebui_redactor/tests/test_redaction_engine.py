from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openwebui_redactor.assessment_redactor_pipe import (
    deterministic_entities,
    extract_docx_text,
    extract_pdf_text,
    redact_docx,
    redact_pdf,
)


SOURCE_TEXT = (
    "Cognitive Process Profile (CPP)\n"
    "Standard Report for Example Organisation\n"
    "STRICTLY CONFIDENTIAL\n"
    "NAME: Alex Taylor    CPP NUMBER: CPP-77821\n"
    "EMAIL: candidate@example.com\n"
    "PHONE: +27 00 000 0000\n"
    "ID NUMBER: TEST-ID-0001\n"
    "ASSESSMENT DATE: 2022-07-25\n"
    "REPORT DATE: 2022-08-01\n"
    "Gender: Female\n"
    "Nationality: South Africa\n"
    "Ethnicity: Black African\n"
    "Highest education: Multiple Degrees / Postgraduate\n"
    "Discipline: Psychology / Social Science\n"
    "Functional area: Human Resources\n"
    "Current position: Trainee\n"
    "Colour blind: No\n"
    "Previous CPP: No\n"
    "How well could you concentrate? Not very well\n\n"
    "Alex Taylor obtained a percentile score of 73.\n"
    "No changes are required to the score wording.\n"
    "The assessment wording and score must remain unchanged.\n"
    "Publisher telephone: +27 11 000 0000"
)


class RedactionEngineTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.candidate_id = "CAND-0042"

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def test_deterministic_detection(self) -> None:
        entities = deterministic_entities(SOURCE_TEXT)
        values = {(entity.text, entity.category) for entity in entities}
        self.assertIn(("Alex Taylor", "CANDIDATE_NAME"), values)
        self.assertIn(("CPP-77821", "CANDIDATE_ID"), values)
        self.assertIn(("candidate@example.com", "EMAIL"), values)
        self.assertIn(("+27 00 000 0000", "PHONE"), values)
        self.assertIn(("TEST-ID-0001", "NATIONAL_ID"), values)
        self.assertNotIn(("+27 11 000 0000", "PHONE"), values)
        self.assertIn(("2022-07-25", "ASSESSMENT_DATE"), values)
        self.assertIn(("2022-08-01", "REPORT_DATE"), values)
        self.assertIn(("Female", "GENDER"), values)
        self.assertIn(("Black African", "ETHNICITY"), values)
        self.assertIn(("Not very well", "SELF_EVALUATION"), values)
        self.assertNotIn(("2022-07-25", "DATE_OF_BIRTH"), values)

    def _make_pdf(self, path: Path) -> None:
        document = fitz.open()
        page = document.new_page(width=595, height=842)
        page.insert_textbox(
            fitz.Rect(55, 55, 540, 790),
            SOURCE_TEXT,
            fontname="helv",
            fontsize=10,
            lineheight=1.4,
        )
        document.set_metadata({"author": "Alex Taylor", "title": "Alex Taylor CPP"})
        document.save(path)
        document.close()

    def test_pdf_redaction_preserves_page_and_scores(self) -> None:
        source = self.root / "source.pdf"
        output = self.root / "output.pdf"
        self._make_pdf(source)
        text, _ = extract_pdf_text(source)
        result = redact_pdf(source, output, deterministic_entities(text), self.candidate_id)

        original = fitz.open(source)
        redacted = fitz.open(output)
        try:
            self.assertEqual(original.page_count, redacted.page_count)
            self.assertEqual(original[0].rect, redacted[0].rect)
            redacted_text = redacted[0].get_text("text")
            self.assertNotIn("Alex Taylor", redacted_text)
            self.assertNotIn("CPP-77821", redacted_text)
            self.assertNotIn("candidate@example.com", redacted_text)
            self.assertIn(self.candidate_id, redacted_text)
            self.assertIn("percentile score of 73", redacted_text)
            self.assertNotIn("2022-07-25", redacted_text)
            self.assertNotIn("2022-08-01", redacted_text)
            self.assertNotIn("Example Organisation", redacted_text)
            self.assertNotIn("Female", redacted_text)
            self.assertNotIn("Black African", redacted_text)
            self.assertNotIn("Not very well", redacted_text)
            self.assertIn("No changes are required", redacted_text)
            self.assertIn("Publisher telephone: +27 11 000 0000", redacted_text)
            self.assertFalse(redacted.metadata.get("author"))
            self.assertGreater(sum(result.counts.values()), 0)
        finally:
            original.close()
            redacted.close()

    def _make_docx(self, path: Path) -> None:
        document = Document()
        section = document.sections[0]
        header = section.header.paragraphs[0]
        header.text = "STRICTLY CONFIDENTIAL — Alex Taylor"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title = document.add_heading("Cognitive Process Profile", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_paragraph("NAME: ")
        paragraph.add_run("Alex ").bold = True
        paragraph.add_run("Taylor").bold = True
        paragraph.add_run("    CPP NUMBER: CPP-77821")

        table = document.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Email"
        table.cell(0, 1).text = "candidate@example.com"
        table.cell(1, 0).text = "Phone"
        table.cell(1, 1).text = "+27 00 000 0000"
        table.cell(2, 0).text = "ID Number"
        table.cell(2, 1).text = "TEST-ID-0001"

        document.add_paragraph("Standard Report for Example Organisation")
        document.add_paragraph("ASSESSMENT DATE: 2022-07-25")
        document.add_paragraph("REPORT DATE: 2022-08-01")

        biography = document.add_table(rows=9, cols=2)
        biography.style = "Table Grid"
        for row, (label, value) in enumerate(
            (
                ("Gender", "Female"),
                ("Nationality", "South Africa"),
                ("Ethnicity", "Black African"),
                ("Highest education", "Multiple Degrees / Postgraduate"),
                ("Discipline", "Psychology / Social Science"),
                ("Functional area", "Human Resources"),
                ("Current position", "Trainee"),
                ("Colour blind", "No"),
                ("Previous CPP", "No"),
            )
        ):
            biography.cell(row, 0).text = label
            biography.cell(row, 1).text = value

        document.add_paragraph("How well could you concentrate? Not very well")
        document.add_paragraph("Alex Taylor obtained a percentile score of 73.")
        document.add_paragraph("No changes are required to the score wording.")
        document.add_paragraph("Publisher telephone: +27 11 000 0000")
        document.core_properties.author = "Alex Taylor"
        document.core_properties.title = "Alex Taylor CPP"
        document.save(path)

    def test_docx_redaction_preserves_structure_and_split_runs(self) -> None:
        source = self.root / "source.docx"
        output = self.root / "output.docx"
        self._make_docx(source)

        text, _ = extract_docx_text(source)
        result = redact_docx(source, output, deterministic_entities(text), self.candidate_id)
        source_doc = Document(source)
        output_doc = Document(output)

        self.assertEqual(len(source_doc.paragraphs), len(output_doc.paragraphs))
        self.assertEqual(len(source_doc.tables), len(output_doc.tables))
        self.assertEqual(source_doc.paragraphs[0].style.name, output_doc.paragraphs[0].style.name)
        self.assertEqual(source_doc.tables[0].style.name, output_doc.tables[0].style.name)

        output_text, _ = extract_docx_text(output)
        self.assertNotIn("Alex Taylor", output_text)
        self.assertNotIn("CPP-77821", output_text)
        self.assertNotIn("candidate@example.com", output_text)
        self.assertIn(self.candidate_id, output_text)
        self.assertIn("percentile score of 73", output_text)
        self.assertNotIn("2022-07-25", output_text)
        self.assertNotIn("2022-08-01", output_text)
        self.assertNotIn("Example Organisation", output_text)
        self.assertNotIn("Female", output_text)
        self.assertNotIn("Black African", output_text)
        self.assertNotIn("Not very well", output_text)
        self.assertIn("No changes are required", output_text)
        self.assertIn("Publisher telephone: +27 11 000 0000", output_text)
        self.assertFalse(output_doc.core_properties.author)
        self.assertGreater(sum(result.counts.values()), 0)

        with zipfile.ZipFile(output, "r") as archive:
            all_xml = b"\n".join(
                archive.read(item)
                for item in archive.infolist()
                if item.filename.endswith(".xml") or item.filename.endswith(".rels")
            )
        self.assertNotIn(b"Alex Taylor", all_xml)
        self.assertNotIn(b"candidate@example.com", all_xml)


if __name__ == "__main__":
    unittest.main()
